# Import necessary libraries
import streamlit as st
import io
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
import pandas as pd

# Create a new Word document
doc = Document()

# Set default font settings
default_font_name = 'Calibri'
default_font_size = Pt(12)
default_paragraph_alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
default_font_color = RGBColor(0, 0, 0)

# Apply styles to the Normal style
normal_style = doc.styles['Normal']
normal_style.font.name = default_font_name
normal_style.font.size = default_font_size
normal_style.paragraph_format.alignment = default_paragraph_alignment
normal_style.font.color.rgb = default_font_color

# Streamlit app title and introduction
st.title("Resume Generator")

# Initialize session state variables
section_keys = ['prof_exp', 'education', 'skills', 'certifications', 'projects', 'prof_memberships', 'languages', 'volunteer']
session_state_keys = [f"{key}_rows" for key in section_keys]

for key in session_state_keys:
    if key not in st.session_state:
        st.session_state[key] = 0

# Function to create a table in Streamlit
def create_table_with_predefined_size(rows, cols, key_prefix):
    data = []
    for row in range(rows):
        row_data = [st.text_input(f"Enter value for row {row + 1}, column {col + 1}:", key=f"{key_prefix}-{row}-{col}") for col in range(cols)]
        data.append(row_data)

    df = pd.DataFrame(data, columns=[f"Column {col + 1}" for col in range(cols)])
    st.table(df)

# Function to get user input for string values
def get_user_input(prompt):
    return st.text_input(prompt).strip()

# Section header (Title)
st.write("We need some basic information to generate your resume.")

# Get user input for personal details
full_name = get_user_input("Enter your full name:")
phone_number = get_user_input("Enter your phone number:")
email_address = get_user_input("Enter your email address:")
linkedin_profile = get_user_input("Enter your LinkedIn profile (if applicable):")
portfolio_website = get_user_input("Enter your professional website or portfolio (if applicable):")

# Page numbers in the footer
for section in doc.sections:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("Page ")
    run.bold = False

# Resume section
st.header("Resume")
st.write("Please enter your professional details.")

# Add user details to the Word document
doc.add_paragraph(f"Full Name: {full_name}")
doc.add_paragraph(f"Phone Number: {phone_number}")
doc.add_paragraph(f"Email Address: {email_address}")
doc.add_paragraph(f"LinkedIn Profile: {linkedin_profile}")
doc.add_paragraph(f"Professional Website/Portfolio: {portfolio_website}")

# Resume Summary or Objective
st.subheader("Resume Summary or Objective")
resume_summary = st.text_area("Enter your resume summary or objective:")
doc.add_heading("Resume Summary or Objective", level=2)
doc.add_paragraph(resume_summary)

# Loop through resume sections and create tables
for i, section_key in enumerate(section_keys):
    st.subheader(section_key.capitalize().replace("_", " "))
    num_rows = st.number_input(f"Enter the number of rows for {section_key.capitalize().replace('_', ' ')}:", min_value=0, value=3, key=f"unique_key_for_{section_key}_rows")
    create_table_with_predefined_size(num_rows, 4, i + 1)  # Specify the number of columns as needed

# Achievements and Awards
st.subheader("Achievements and Awards")
achievements = st.text_area("Enter your achievements and awards:")

# Publications/Presentations
st.subheader("Publications/Presentations")
publications = st.text_area("Enter your publications or presentations:")

# Hobbies and Interests
st.subheader("Hobbies and Interests")
hobbies_interests = st.text_area("Enter your hobbies and interests:")

# References
st.subheader("References")
references = st.text_area("Enter your references information:")

# Save the Word document with cover letter content
file_name = st.text_input("Enter the file name for saving the resume")

if file_name:
    file_path = f"{file_name}.docx"

    # Download button
    def main():
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        st.download_button("Download DOCX", data=output.getvalue(), file_name=file_name)

    main()

st.write("Resume generator completed.")
