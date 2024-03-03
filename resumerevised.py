import streamlit as st
import io
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
import pandas as pd
import os

# Set up default font settings
default_font_name = 'Calibri'
default_font_size = Pt(12)

# Create a new Word document
doc = Document()

# Apply default font to the Normal style
normal_style = doc.styles['Normal']
normal_style.font.name = default_font_name
normal_style.font.size = default_font_size
normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Streamlit app title and introduction
st.title("AI-Powered Resume and Cover Letter Generator")

# Font selection dropdown
selected_font = st.selectbox("Select font for the document:",
                             options=["Calibri", "Arial", "Times New Roman", "Verdana", "Georgia"])

# Section selection checkboxes
available_sections = ["Summary or Objective", "Professional Experience", "Education",
                      "Skills", "Certifications and Training", "Achievements and Awards",
                      "Projects", "Publications/Presentations", "Professional Memberships",
                      "Languages", "Volunteer Work", "Hobbies and Interests", "References",
                      "Cover Letter"]
selected_sections = st.multiselect("Select sections to include:", available_sections)

# ... (Rest of the code with enhancements for content structure, column names, experience detail, table borders, etc.)

# Cover letter template
def add_cover_letter():
    doc.add_page_break()
    doc.add_heading("Cover Letter", level=1)

    # Add basic cover letter content here, using placeholders for personalization
    doc.add_paragraph(f"Dear [Hiring Manager Name],")
    doc.add_paragraph("I am writing to express my enthusiastic interest in the [Job Title] position at [Company Name].")
    # ... (Continue with cover letter content as needed)

# Generate and save the document
def generate_and_save():
    # Apply selected font to the entire document
    for paragraph in doc.paragraphs:
        paragraph.style.font.name = selected_font
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style.font.name = selected_font

    # Add selected sections to the document
    for section in selected_sections:
        if section == "Summary or Objective":
        # ... (Add summary or objective content)
        elif section == "Professional Experience":
        # ... (Add professional experience content)
        # ... (Add other sections accordingly)
        elif section == "Cover Letter":
            add_cover_letter()

    # Save the document
    file_name = st.text_input("Enter the file name for saving the resume and cover letter:")
    if file_name:
        file_path = f"{file_name}.docx"
        doc.save(file_path)

        # Download button
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        st.download_button("Download DOCX", data=output.getvalue(), file_name=file_name)

        st.write("Resume and Cover Letter generator completed.")

# Call the generation function if a file name is provided
if st.button("Generate"):
    generate_and_save()